"""Build evidence-derived data for the DUCAR app.

The app should communicate what the manuals, TORs, budget reports and global
framework references say, rather than listing sources as inert bibliography.
This script extracts text from the core local documents, folds in the full
manual repository catalogue, checks online source availability, and emits a
compact JSON dataset for charts and tables.
"""

from __future__ import annotations

import io
import json
import re
import statistics
import zipfile
import csv
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from html import unescape
from html.parser import HTMLParser
from pathlib import Path
from urllib.error import URLError
from urllib.request import Request, urlopen

import pandas as pd
from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
PUBLIC_DOCS = ROOT / "public" / "docs"
PUBLIC_DATA = ROOT / "public" / "data"
DATA_DIR = ROOT / "data"
TOR_ROOT = Path(r"D:\OneDrive\Procurements\TOR - DUCACR")
MANUALS_CATALOG = PUBLIC_DATA / "manuals_catalog.json"
MOWT_CATALOG = PUBLIC_DATA / "mowt_manuals_catalog.json"
OUT_PUBLIC = PUBLIC_DATA / "evidence_synthesis.json"
OUT_DATA = DATA_DIR / "evidence_synthesis.json"
CASE_PACKAGE_WORKBOOK = TOR_ROOT / "DUCAR_Framework_Tool" / "evidence_and_case_studies" / "DUCAR_APA_References_and_Assumptions_Register.xlsx"
TRANSPORT_VEHICLE_WORKBOOK = TOR_ROOT / "Road transport data" / "Motor Vehicle data Updated Up to December 2023(1).xls"

SUPPORTED_LOCAL_EXTENSIONS = {".pdf", ".docx", ".json", ".csv", ".xlsx", ".xls", ".md", ".txt"}
EXCLUDED_DIR_NAMES = {
    ".git",
    "__pycache__",
    "node_modules",
    "dist",
    "build",
    ".vite",
}
EXCLUDED_DIR_PREFIXES = ("rendered_",)
EXCLUDED_FILE_NAMES = {
    "package.json",
    "package-lock.json",
    "tsconfig.json",
    "tsconfig.app.json",
    "tsconfig.node.json",
    "_manifest_write_test.json",
}

TOPIC_KEYWORDS = {
    "PIMS appraisal": ["pims", "public investment", "appraisal", "feasibility", "npv", "eirr", "economic", "cost-benefit"],
    "Budget monitoring": ["budget", "release", "expenditure", "absorption", "variance", "fy", "allocation", "financing"],
    "Road condition": ["condition", "roughness", "iri", "rutting", "cracking", "pothole", "visual inspection", "deterioration"],
    "RAM and lifecycle": ["asset management", "rams", "lifecycle", "life cycle", "maintenance strategy", "work plan", "dtims", "hdm"],
    "Traffic and axle load": ["traffic", "aadt", "vehicle", "axle", "overload", "weighbridge", "flow", "speed"],
    "GIS and network": ["gis", "geo", "spatial", "route", "network", "linear referencing", "coordinate", "map"],
    "Construction QA": ["construction", "quality", "materials", "asphalt", "concrete", "specification", "testing", "supervision"],
    "Climate and drainage": ["climate", "rainfall", "flood", "drainage", "culvert", "erosion", "slope", "resilience"],
    "Safety and NMT": ["safety", "crash", "accident", "pedestrian", "cycling", "non-motorized", "crossing", "speed"],
    "Bridge and structures": ["bridge", "structure", "culvert", "bms", "span", "bearing", "deck", "inspection"],
    "Procurement and contracts": ["procurement", "contract", "tender", "payment", "boq", "certificate", "ppda", "contractor"],
}

ONLINE_SOURCES = [
    {
        "group": "National planning and budget",
        "title": "Uganda Vision 2040",
        "agency": "National Planning Authority",
        "url": "https://npa.go.ug/uganda-vision-2040/",
        "apa": "National Planning Authority. (n.d.). Uganda Vision 2040. The Republic of Uganda.",
    },
    {
        "group": "National planning and budget",
        "title": "Budget Speech FY 2024/25",
        "agency": "Ministry of Finance, Planning and Economic Development",
        "url": "https://www.budget.finance.go.ug/content/budget-speech-12",
        "apa": "Ministry of Finance, Planning and Economic Development. (2024). Budget Speech FY 2024/25. The Republic of Uganda.",
    },
    {
        "group": "National planning and budget",
        "title": "Budget Speech FY 2025/26",
        "agency": "Ministry of Finance, Planning and Economic Development",
        "url": "https://budget.finance.go.ug/sites/default/files/National%20Budget%20docs/Budget%20Speech%20FY2025-26.pdf",
        "apa": "Ministry of Finance, Planning and Economic Development. (2025). Budget Speech Financial Year 2025/26. The Republic of Uganda.",
    },
    {
        "group": "National planning and budget",
        "title": "Background to the Budget FY 2024/25",
        "agency": "Ministry of Finance, Planning and Economic Development",
        "url": "https://budget.finance.go.ug/content/background-budget-10",
        "apa": "Ministry of Finance, Planning and Economic Development. (2024). Background to the Budget FY 2024/25. The Republic of Uganda.",
    },
    {
        "group": "National planning and budget",
        "title": "Approved Budget Estimates FY 2024/25",
        "agency": "Ministry of Finance, Planning and Economic Development",
        "url": "https://budget.finance.go.ug/content/approved-budget-estimates-1241",
        "apa": "Ministry of Finance, Planning and Economic Development. (2024). Approved budget estimates, central governments FY 2024/25. The Republic of Uganda.",
    },
    {
        "group": "National planning and budget",
        "title": "National Budget Framework Paper FY 2024/25",
        "agency": "Ministry of Finance, Planning and Economic Development",
        "url": "https://www.budget.finance.go.ug/content/national-budget-framework-paper-15",
        "apa": "Ministry of Finance, Planning and Economic Development. (2024). National budget framework paper FY 2024/25. The Republic of Uganda.",
    },
    {
        "group": "Traffic and geospatial",
        "title": "Uganda Roads",
        "agency": "World Bank Transport Data",
        "url": "https://datacatalog.worldbank.org/infrastructure-data/search/dataset/0041482/Uganda-Roads",
        "apa": "World Bank. (n.d.). Uganda roads. World Bank Transport Data.",
    },
    {
        "group": "Traffic and geospatial",
        "title": "Uganda Road Network main roads",
        "agency": "AmeriGEOSS / WFP OpenStreetMap extract",
        "url": "https://data.amerigeoss.org/dataset/uganda-road-network-main-roads",
        "apa": "AmeriGEOSS. (n.d.). Uganda road network main roads.",
    },
    {
        "group": "Global RAM",
        "title": "PIARC Road Asset Management Manual case studies",
        "agency": "World Road Association",
        "url": "https://road-asset.piarc.org/en/management-asset-management-implementation/case-studies",
        "apa": "World Road Association. (n.d.). Road asset management manual case studies.",
    },
    {
        "group": "Global RAM",
        "title": "FHWA Asset Management",
        "agency": "Federal Highway Administration",
        "url": "https://www.fhwa.dot.gov/asset/",
        "apa": "Federal Highway Administration. (n.d.). Asset management.",
    },
    {
        "group": "Global RAM",
        "title": "Austroads Guide to Asset Management",
        "agency": "Austroads",
        "url": "https://austroads.gov.au/infrastructure/asset-management/guide-to-asset-management",
        "apa": "Austroads. (n.d.). Guide to asset management.",
    },
    {
        "group": "Global RAM",
        "title": "AfDB Road Asset Management Toolkit",
        "agency": "African Development Bank",
        "url": "https://www.afdb.org/en/documents/road-asset-management-study-accelerating-road-sector-reforms-part-ii-road-asset-management-toolkit",
        "apa": "African Development Bank. (n.d.). Road asset management study: Road asset management toolkit.",
    },
    {
        "group": "Global RAM",
        "title": "World Bank performance-based maintenance study",
        "agency": "World Bank",
        "url": "https://documents1.worldbank.org/curated/en/413451468336612648/pdf/689620ESW0P102002012000Final0Report.pdf",
        "apa": "World Bank. (2012). Performance-based contracting for preservation and improvement of road assets.",
    },
    {
        "group": "Global RAM",
        "title": "World Bank PBC review",
        "agency": "World Bank",
        "url": "https://blogs.worldbank.org/en/transport/performance-based-contracts-promoting-quality-road-maintenance-and-economic-efficiency",
        "apa": "World Bank. (n.d.). Performance-based contracts: Promoting quality road maintenance and economic efficiency.",
    },
    {
        "group": "Global RAM",
        "title": "ReCAP effective road asset management baseline",
        "agency": "GOV.UK / ReCAP",
        "url": "https://www.gov.uk/research-for-development-outputs/economic-growth-through-effective-road-asset-management-consolidated-baseline-study-report",
        "apa": "GOV.UK. (n.d.). Economic growth through effective road asset management consolidated baseline study report.",
    },
]

ITIS_TABLES = {
    "road_network_by_category": {
        "title": "Uganda road network by category, FY 2022/23",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; URF DUCAR network status report 2023.",
        "columns": ["Category", "Length_km", "DUCAR_scope"],
        "rows": [
            ["National Roads", 21292, "Reference only"],
            ["District Roads", 41194, "DUCAR"],
            ["KCCA", 2103, "DUCAR urban"],
            ["City Roads", 2830, "DUCAR urban"],
            ["Community Access Roads", 75404, "DUCAR"],
            ["Town Council Roads", 24269, "DUCAR urban"],
            ["Municipal Roads", 6656, "DUCAR urban"],
        ],
    },
    "road_condition_by_category": {
        "title": "Road condition by network category, FY 2022/23",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; URF DUCAR network status report 2023.",
        "columns": ["Network category", "Good_km", "Fair_km", "Poor_km", "Total_km"],
        "rows": [
            ["National (UNRA)", 12508, 7844, 848, 21200],
            ["KCCA", 184.9, 1019.5, 898.6, 2103],
            ["Community Access Roads", 3588.45, 1159.68, 70656, 75404],
            ["City Roads", 529.64, 200.14, 2099.83, 2830],
            ["District Roads", 4008.69, 1865.43, 35320, 41194],
            ["Municipal Roads", 505.09, 410.7, 5740.5, 6656],
            ["Town Council Roads", 673, 480.56, 23115, 24269],
        ],
    },
    "paved_national_roads_trend": {
        "title": "Paved national road network trend",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; UNRA annual performance report 2022/23.",
        "columns": ["FY", "Annual increase_km", "Paved stock_km", "Percent paved"],
        "rows": [
            ["2018/19", 420, 4942, 23.5],
            ["2019/20", 428, 5370, 25.5],
            ["2020/21", 221, 5591, 26.6],
            ["2021/22", 287, 5878.5, 27.8],
            ["2022/23", 254.5, 6133, 29.18],
        ],
    },
    "road_crashes_by_nature": {
        "title": "Road traffic crashes by nature, CY 2019-2023",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; Uganda Police Force annual crime report 2023.",
        "columns": ["Year", "Fatal", "Serious", "Minor", "Total"],
        "rows": [
            [2019, 3407, 5992, 3459, 12858],
            [2020, 3269, 5803, 3177, 12249],
            [2021, 3757, 9070, 4616, 17443],
            [2022, 3901, 10776, 5717, 20394],
            [2023, 4179, 12487, 6942, 23608],
        ],
    },
    "rail_network": {
        "title": "Rail network operational status, 2023",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; URC annual performance report 2022/23.",
        "columns": ["Status", "Km", "Share_percent"],
        "rows": [["Operational", 269, 21.2], ["Non-operational", 997, 78.8]],
    },
    "railway_accidents": {
        "title": "Railway accidents by category",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; URC annual performance report 2022/23.",
        "columns": ["Category", "FY 2018/19", "FY 2019/20", "FY 2020/21", "FY 2021/22", "FY 2022/23"],
        "rows": [["Fatal", 0, 0, 1, 2, 2], ["Serious", 0, 0, 0, 1, 1], ["Minor", 22, 19, 48, 30, 26]],
    },
    "air_passenger_and_cargo": {
        "title": "SDG 9.1.2 passenger and freight indicators",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023.",
        "columns": ["Mode", "Measure", "2019", "2020", "2021", "2022", "2023"],
        "rows": [
            ["Air transport", "Cargo tonnes", 64731, 59720, 64172, 61255, 59072],
            ["Air transport", "Passengers", 1802107, 565541, 941688, 1574405, 1932094],
            ["Railway transport", "Cargo tonnes", 193693, 196935, 390628, 242686, 243634.18],
            ["Railway transport", "Passengers", 619206, 180327, 799533, 155816, 216455],
            ["Water transport", "Cargo tonnes", 1845.6, 2110.1, 2834.7, None, 96922],
            ["Water transport", "Passengers", 1155299, 3842092, 1623852, 3117791, 3518023],
        ],
    },
    "ferry_operations_kis": {
        "title": "KIS ferry crossings and traffic, CY 2019-2023",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; KIS quarterly independent operations reports.",
        "columns": ["Year", "Passengers", "Pick-ups", "Trucks", "Cars", "Buses", "Motor bike", "Total traffic"],
        "rows": [
            [2019, 547104, 7016, 11714, 18934, 5337, 42157, 632262],
            [2020, 601815, 7718, 12886, 20827, 5870, 46373, 695489],
            [2021, 661997, 8489, 14174, 22910, 6457, 51010, 765037],
            [2022, 728375, 8923, 19012, 39227, 8505, 83721, 887763],
            [2023, 770925, 9462, 18496, 41539, 8706, 77588, 926716],
        ],
    },
    "water_body_accidents": {
        "title": "Recorded accidents on major water bodies",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; Uganda Police Force annual crime report 2023.",
        "columns": ["Year", "Passengers involved", "Passengers rescued", "Fatalities"],
        "rows": [[2019, 1503, 1214, 289], [2020, 0, 0, 0], [2021, 456, 246, 210], [2022, 291, 67, 224], [2023, 566, 323, 243]],
    },
    "major_lakes": {
        "title": "Major water bodies and transport setting",
        "source": "Ministry of Works and Transport. (2023). ITIS statistical abstract 2023; ResearchGate lake characteristics.",
        "columns": ["Lake", "Shoreline_km", "Area_km2"],
        "rows": [["Victoria", 4828, 68800], ["Albert", 355, 2675], ["Kyoga", 200, 1720], ["Edward", 234, 2325], ["George", 63.2, 250], ["Bisina", None, 308]],
    },
}


class TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        clean = data.strip()
        if clean:
            self.parts.append(clean)


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", unescape(text or "")).strip()


def count_words(text: str) -> int:
    return len(re.findall(r"[A-Za-z][A-Za-z0-9'-]*", text))


def topic_counts(text: str) -> dict[str, int]:
    lower = text.lower()
    scores: dict[str, int] = {}
    for topic, terms in TOPIC_KEYWORDS.items():
        scores[topic] = sum(len(re.findall(rf"\b{re.escape(term.lower())}\b", lower)) for term in terms)
    return scores


def top_sentences(text: str, limit: int = 4) -> list[str]:
    candidates = re.split(r"(?<=[.!?])\s+", text)
    scored = []
    priority = ["road", "budget", "asset", "condition", "traffic", "maintenance", "investment", "ducar", "pims", "bridge", "safety"]
    for sentence in candidates:
        sentence = sentence.strip()
        if len(sentence) < 80 or len(sentence) > 360:
            continue
        lower = sentence.lower()
        score = sum(lower.count(term) for term in priority) + min(5, len(re.findall(r"\d", sentence)))
        if score:
            scored.append((score, sentence))
    return [sentence for _, sentence in sorted(scored, reverse=True)[:limit]]


def extract_pdf(path: Path) -> tuple[str, int, int]:
    reader = PdfReader(str(path))
    pages = len(reader.pages)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return normalize_text(" ".join(parts)), pages, 0


def extract_docx(path: Path) -> tuple[str, int, int]:
    doc = Document(str(path))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    table_count = len(doc.tables)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return normalize_text(" ".join(parts)), 0, table_count


def serialise_cell(value) -> str | int | float | None:
    if pd.isna(value):
        return None
    if isinstance(value, (int, float)):
        if isinstance(value, float) and value.is_integer():
            return int(value)
        return round(float(value), 4)
    text = str(value).strip()
    if not text:
        return None
    if len(text) > 220:
        return f"{text[:217]}..."
    return text


def table_preview_from_frame(df: pd.DataFrame, title: str, source: str, max_rows: int = 8, max_cols: int = 8) -> dict:
    clean = df.dropna(how="all").dropna(axis=1, how="all")
    if clean.empty:
        return {"title": title, "source": source, "columns": [], "rows": [], "row_count": 0, "column_count": 0}

    preview = clean.iloc[:max_rows, :max_cols]
    rows = [[serialise_cell(value) for value in row] for row in preview.to_numpy().tolist()]
    columns = [f"Column {index + 1}" for index in range(preview.shape[1])]
    return {
        "title": title,
        "source": source,
        "columns": columns,
        "rows": rows,
        "row_count": int(clean.shape[0]),
        "column_count": int(clean.shape[1]),
    }


def extract_workbook(path: Path) -> tuple[str, int, int, list[dict]]:
    xl = pd.ExcelFile(path)
    parts: list[str] = []
    previews: list[dict] = []
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name, header=None, dtype=object)
        clean = df.dropna(how="all").dropna(axis=1, how="all")
        if clean.empty:
            continue
        cells = [str(value) for value in clean.to_numpy().ravel() if not pd.isna(value)]
        parts.extend(cells[:5000])
        previews.append(table_preview_from_frame(clean, f"{path.stem} / {sheet_name}", str(path)))
    return normalize_text(" ".join(parts)), 0, len(previews), previews


def extract_csv(path: Path) -> tuple[str, int, int, list[dict]]:
    rows: list[list[str]] = []
    with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            rows.append(row)
    text = normalize_text(" ".join(" ".join(cell for cell in row if cell) for row in rows))
    if not rows:
        return text, 0, 0, []
    max_cols = max(len(row) for row in rows)
    frame = pd.DataFrame([row + [""] * (max_cols - len(row)) for row in rows])
    return text, 0, 1, [table_preview_from_frame(frame, path.stem, str(path))]


def extract_plain_text(path: Path) -> tuple[str, int, int, list[dict]]:
    return normalize_text(path.read_text(encoding="utf-8", errors="ignore")), 0, 0, []


def source_area(path: Path) -> str:
    text = path.as_posix().lower()
    if "/evidence_and_case_studies/" in text:
        return "Global case evidence package"
    if "/road transport data/" in text:
        return "Road transport data"
    if "/workbooks/" in text:
        return "Generated DUCAR workbooks"
    if "/manuals/" in text or "/policies/" in text:
        return "DUCAR manuals and policies"
    if "/public/docs/" in text or "/data_sources/" in text:
        return "Published source documents"
    if "/github_app/data/" in text or "/public/data/" in text:
        return "Web data feeds"
    if "/gis/" in text or "/roads/" in text or "/districts/" in text:
        return "GIS and road data"
    return "TOR and local source documents"


def read_doc(path: Path) -> dict:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            text, pages, tables = extract_pdf(path)
            table_extracts = []
        elif ext == ".docx":
            text, pages, tables = extract_docx(path)
            table_extracts = []
        elif ext in {".xlsx", ".xls"}:
            text, pages, tables, table_extracts = extract_workbook(path)
        elif ext == ".csv":
            text, pages, tables, table_extracts = extract_csv(path)
        elif ext in {".md", ".txt"}:
            text, pages, tables, table_extracts = extract_plain_text(path)
        elif ext == ".json":
            text = normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
            pages, tables, table_extracts = 0, 0, []
        else:
            text, pages, tables, table_extracts = "", 0, 0, []
        counts = topic_counts(text)
        dominant = max(counts.items(), key=lambda item: item[1])[0] if text else "Not text-extracted"
        try:
            relative = path.relative_to(TOR_ROOT).as_posix()
        except ValueError:
            relative = path.name
        return {
            "name": path.name,
            "path": relative,
            "source_area": source_area(path),
            "extension": ext,
            "bytes": path.stat().st_size,
            "pages": pages,
            "tables": tables,
            "table_extracts": table_extracts[:3],
            "words": count_words(text),
            "characters": len(text),
            "dominant_topic": dominant,
            "topic_counts": counts,
            "summary_points": top_sentences(text),
            "status": "read" if text else "metadata only",
        }
    except Exception as exc:
        return {
            "name": path.name,
            "path": path.name,
            "source_area": source_area(path),
            "extension": ext,
            "bytes": path.stat().st_size if path.exists() else 0,
            "pages": 0,
            "tables": 0,
            "table_extracts": [],
            "words": 0,
            "characters": 0,
            "dominant_topic": "Extraction issue",
            "topic_counts": {topic: 0 for topic in TOPIC_KEYWORDS},
            "summary_points": [],
            "status": f"error: {type(exc).__name__}",
        }


def iter_core_documents() -> list[Path]:
    candidates: dict[str, Path] = {}
    if not TOR_ROOT.exists():
        return []
    for path in TOR_ROOT.rglob("*"):
        if not path.is_file() or path.name.startswith("~$"):
            continue
        if path.suffix.lower() not in SUPPORTED_LOCAL_EXTENSIONS:
            continue
        if path.name.lower() in EXCLUDED_FILE_NAMES:
            continue
        parts = {part.lower() for part in path.parts}
        if parts & EXCLUDED_DIR_NAMES:
            continue
        if any(part.lower().startswith(EXCLUDED_DIR_PREFIXES) for part in path.parts):
            continue
        candidates[path.resolve().as_posix().lower()] = path
    return sorted(candidates.values(), key=lambda item: item.name.lower())


def load_json(path: Path, fallback):
    if not path.exists():
        return fallback
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return fallback


def read_manual_repository_catalog() -> dict:
    catalog = load_json(MANUALS_CATALOG, {})
    summary = catalog.get("summary", {})
    files = catalog.get("files", [])
    logic_records = catalog.get("logic_records", [])
    by_topic = catalog.get("by_topic", {})
    by_role = catalog.get("by_evidence_role", {})
    by_folder = catalog.get("by_folder", {})
    by_extension = catalog.get("by_extension", {})
    return {
        "summary": summary,
        "by_topic": by_topic,
        "by_role": by_role,
        "by_folder_top": dict(list(by_folder.items())[:12]),
        "by_extension_top": dict(list(by_extension.items())[:12]),
        "logic_record_count": summary.get("logic_records", len(logic_records)),
        "file_count": summary.get("all_files", len(files)),
    }


def fetch_online_source(source: dict) -> dict:
    url = source["url"]
    result = {**source, "status": "unread", "bytes": 0, "words": 0, "topic_counts": {topic: 0 for topic in TOPIC_KEYWORDS}, "evidence_score": 0}
    try:
        req = Request(url, headers={"User-Agent": "DUCAR Priority Studio evidence reader/1.0"})
        with urlopen(req, timeout=18) as response:
            raw = response.read(1_500_000)
            content_type = response.headers.get("content-type", "")
        result["bytes"] = len(raw)
        if "pdf" in content_type.lower() or url.lower().endswith(".pdf"):
            try:
                reader = PdfReader(io.BytesIO(raw))
                text = normalize_text(" ".join((page.extract_text() or "") for page in reader.pages[:30]))
                result["pages_sampled"] = min(len(reader.pages), 30)
                result["pages_reported"] = len(reader.pages)
            except Exception:
                text = ""
        else:
            parser = TextExtractor()
            parser.feed(raw.decode("utf-8", errors="ignore"))
            text = normalize_text(" ".join(parser.parts))
        counts = topic_counts(text)
        result.update({
            "status": "read" if text else "fetched",
            "words": count_words(text),
            "topic_counts": counts,
            "dominant_topic": max(counts.items(), key=lambda item: item[1])[0] if counts else "No text",
            "evidence_score": min(100, 35 + count_words(text) // 120 + sum(1 for value in counts.values() if value > 0) * 5),
            "summary_points": top_sentences(text, 2),
        })
    except (URLError, TimeoutError, OSError, ValueError) as exc:
        result["status"] = f"unavailable: {type(exc).__name__}"
    return result


def build_document_topic_chart(documents: list[dict]) -> list[dict]:
    totals = Counter()
    for doc in documents:
        totals.update(doc.get("topic_counts", {}))
    max_value = max(totals.values() or [1])
    return [
        {
            "topic": topic,
            "mentions": count,
            "share": round((count / max_value) * 100, 1),
            "decision_use": {
                "PIMS appraisal": "Project admission, appraisal readiness, affordability and gate control.",
                "Budget monitoring": "Budget release, absorption, variance and physical progress checks.",
                "Road condition": "Condition-based treatment triggers, deterioration pressure and monitoring tiers.",
                "RAM and lifecycle": "Lifecycle planning, work standards, system governance and optimisation.",
                "Traffic and axle load": "AADT, representative vehicles, axle-load pressure and speed-flow assumptions.",
                "GIS and network": "Route identity, district joins, topology, route matrix and spatial equity.",
                "Construction QA": "Design, material, testing, supervision and quality acceptance rules.",
                "Climate and drainage": "Flood, drainage, culvert and climate-resilience screens.",
                "Safety and NMT": "Crash exposure, vulnerable users and safety countermeasure benefits.",
                "Bridge and structures": "Bridge/culvert inventory, inspection and structural criticality gates.",
                "Procurement and contracts": "Tender readiness, BOQ, contract administration and payment controls.",
            }[topic],
        }
        for topic, count in totals.most_common()
    ]


def build_document_table(documents: list[dict]) -> dict:
    rows = []
    for doc in sorted(documents, key=lambda item: item["words"], reverse=True):
        rows.append([
            doc["name"],
            doc.get("source_area", "Local source"),
            doc["extension"].replace(".", "").upper(),
            doc["pages"] or "",
            doc["tables"] or "",
            doc["words"],
            doc["dominant_topic"],
            doc["status"],
        ])
    return {
        "title": "Local evidence files extracted from TOR - DUCACR",
        "columns": ["Document", "Source area", "Type", "Pages", "Tables", "Words read", "Dominant topic", "Read status"],
        "rows": rows,
    }


def simple_chart(title: str, columns: list[str], counter: Counter | dict) -> dict:
    rows = [[key, value] for key, value in sorted(counter.items(), key=lambda item: item[1], reverse=True)]
    return {"title": title, "columns": columns, "rows": rows}


def build_source_coverage(documents: list[dict]) -> dict:
    by_type = Counter(doc["extension"].replace(".", "").upper() or "No extension" for doc in documents)
    by_area = Counter(doc.get("source_area", "Local source") for doc in documents)
    by_status = Counter(doc.get("status", "unknown") for doc in documents)
    by_topic = Counter(doc.get("dominant_topic", "Unclassified") for doc in documents)
    return {
        "fileTypeChart": simple_chart("Local evidence files by type", ["Type", "Files"], by_type),
        "sourceAreaChart": simple_chart("Local evidence files by source area", ["Source area", "Files"], by_area),
        "extractionStatusChart": simple_chart("Extraction status", ["Status", "Files"], by_status),
        "dominantTopicChart": simple_chart("Dominant topics by file", ["Topic", "Files"], by_topic),
    }


def build_tabular_extracts(documents: list[dict]) -> list[dict]:
    extracts: list[dict] = []
    for doc in documents:
        for table in doc.get("table_extracts", []):
            if table.get("rows"):
                extracts.append(
                    {
                        **table,
                        "document": doc["name"],
                        "source_area": doc.get("source_area", "Local source"),
                        "source": doc.get("path", doc["name"]),
                    }
                )
    return sorted(extracts, key=lambda item: (item.get("source_area", ""), item.get("title", "")))


def dataframe_table(title: str, df: pd.DataFrame, source: str, max_rows: int = 80) -> dict:
    clean = df.dropna(how="all").dropna(axis=1, how="all")
    columns = [str(column) for column in clean.columns]
    rows = [
        [serialise_cell(value) for value in row]
        for row in clean.head(max_rows).to_numpy().tolist()
    ]
    return {
        "title": title,
        "source": source,
        "columns": columns,
        "rows": rows,
        "row_count": int(clean.shape[0]),
        "column_count": int(clean.shape[1]),
    }


def read_case_package_tables() -> dict:
    if not CASE_PACKAGE_WORKBOOK.exists():
        return {}
    tables: dict[str, dict] = {}
    for sheet in ["APA_References", "Decision_Assumptions", "Country_Case_Studies"]:
        df = pd.read_excel(CASE_PACKAGE_WORKBOOK, sheet_name=sheet)
        key = {
            "APA_References": "apaReferences",
            "Decision_Assumptions": "decisionAssumptions",
            "Country_Case_Studies": "countryCaseStudies",
        }[sheet]
        tables[key] = dataframe_table(f"Global case package: {sheet.replace('_', ' ')}", df, CASE_PACKAGE_WORKBOOK.name)
    return tables


def build_global_case_charts(case_tables: dict) -> dict:
    country_table = case_tables.get("countryCaseStudies", {})
    reference_table = case_tables.get("apaReferences", {})
    country_rows = country_table.get("rows", [])
    reference_rows = reference_table.get("rows", [])
    continent_counts = Counter(row[0] for row in country_rows if row and row[0])
    source_counts = Counter(row[-1] for row in country_rows if row and row[-1])
    reference_type_counts = Counter(row[1] for row in reference_rows if len(row) > 1 and row[1])
    return {
        "continentChart": simple_chart("Country case studies by continent", ["Continent", "Cases"], continent_counts),
        "sourceKeyChart": simple_chart("Case-study source keys", ["Source key", "Cases"], source_counts),
        "referenceTypeChart": simple_chart("APA register by source type", ["Source type", "References"], reference_type_counts),
    }


def build_transport_vehicle_charts() -> dict:
    if not TRANSPORT_VEHICLE_WORKBOOK.exists():
        return {}
    df = pd.read_excel(TRANSPORT_VEHICLE_WORKBOOK, sheet_name=0, header=None)
    if df.empty or df.shape[0] < 3:
        return {}
    date_row = df.iloc[1]
    date_columns: list[tuple[int, int]] = []
    for column_index, value in enumerate(date_row):
        parsed = pd.to_datetime(value, errors="coerce")
        if pd.notna(parsed):
            date_columns.append((column_index, int(parsed.year)))

    class_totals: Counter[str] = Counter()
    annual_totals: Counter[int] = Counter()
    class_year: defaultdict[str, Counter[int]] = defaultdict(Counter)
    for _, row in df.iloc[2:].iterrows():
        vehicle_class = str(row.iloc[1]).strip() if len(row) > 1 and not pd.isna(row.iloc[1]) else ""
        if not vehicle_class or vehicle_class.lower() in {"nan", "total"}:
            continue
        for column_index, year in date_columns:
            value = pd.to_numeric(row.iloc[column_index], errors="coerce")
            if pd.isna(value):
                continue
            amount = float(value)
            class_totals[vehicle_class] += amount
            annual_totals[year] += amount
            class_year[vehicle_class][year] += amount

    top_classes = class_totals.most_common(12)
    latest_year = max(annual_totals) if annual_totals else None
    latest_mix = []
    if latest_year:
        latest_mix = sorted(
            [[vehicle_class, round(values[latest_year], 1)] for vehicle_class, values in class_year.items() if values[latest_year]],
            key=lambda row: row[1],
            reverse=True,
        )[:12]
    return {
        "vehicleClassTotals": {
            "title": "Motor vehicle registrations by class, July 2018-February 2024",
            "source": TRANSPORT_VEHICLE_WORKBOOK.name,
            "columns": ["Vehicle class", "Registrations"],
            "rows": [[name, round(total, 1)] for name, total in top_classes],
        },
        "annualVehicleTotals": {
            "title": "Motor vehicle registrations by calendar year",
            "source": TRANSPORT_VEHICLE_WORKBOOK.name,
            "columns": ["Year", "Registrations"],
            "rows": [[year, round(total, 1)] for year, total in sorted(annual_totals.items())],
        },
        "latestYearClassMix": {
            "title": f"Motor vehicle class mix in {latest_year}" if latest_year else "Motor vehicle class mix",
            "source": TRANSPORT_VEHICLE_WORKBOOK.name,
            "columns": ["Vehicle class", "Registrations"],
            "rows": latest_mix,
        },
    }


def build_story_cards(summary: dict, topic_chart: list[dict], case_tables: dict, transport_charts: dict, source_coverage: dict) -> list[dict]:
    top_topic = topic_chart[0] if topic_chart else {"topic": "Evidence", "mentions": 0, "decision_use": "Source material indexed."}
    global_case_count = (case_tables.get("countryCaseStudies") or {}).get("row_count", 0)
    decision_count = (case_tables.get("decisionAssumptions") or {}).get("row_count", 0)
    transport_total = sum(float(row[1] or 0) for row in (transport_charts.get("annualVehicleTotals") or {}).get("rows", [])) if transport_charts else 0
    file_type_rows = (source_coverage.get("fileTypeChart") or {}).get("rows", [])
    leading_type = file_type_rows[0][0] if file_type_rows else "Files"
    return [
        {
            "title": "Local evidence corpus",
            "metric": f"{summary.get('core_documents_found', 0):,}",
            "label": "files indexed",
            "story": "TORs, manuals, policy notes, budget reports, web data feeds, spreadsheets and global case package files are now read into one evidence dataset.",
            "evidence": f"{summary.get('core_documents_read', 0):,} files yielded text or tables; leading type: {leading_type}.",
            "tone": "blue",
        },
        {
            "title": "Decision-topic spine",
            "metric": f"{int(top_topic.get('mentions', 0)):,}",
            "label": top_topic.get("topic", "Topic"),
            "story": top_topic.get("decision_use", "The dominant topic drives decision charts and screening logic."),
            "evidence": "Keyword counts are calculated across the extracted local evidence text.",
            "tone": "green",
        },
        {
            "title": "Global case transfer",
            "metric": f"{global_case_count:,}",
            "label": "case rows",
            "story": "The international case package translates country practices into DUCAR-specific lessons and adaptation rules.",
            "evidence": f"{decision_count:,} decision assumptions are linked back to APA source keys.",
            "tone": "gold",
        },
        {
            "title": "Transport demand signal",
            "metric": f"{round(transport_total):,}",
            "label": "vehicle records",
            "story": "The local motor-vehicle workbook is converted into class and annual demand charts for traffic-pressure context.",
            "evidence": "Values are aggregated from monthly vehicle registration columns in the Road transport data folder.",
            "tone": "cyan",
        },
        {
            "title": "Manual repository depth",
            "metric": f"{summary.get('manual_repository_files', 0):,}",
            "label": "manual files",
            "story": "The national manual repository remains indexed as a large evidence backdrop for RAM, GIS, bridge, condition and QA rules.",
            "evidence": f"{summary.get('manual_logic_records', 0):,} logic records are available for evidence-role mapping.",
            "tone": "red",
        },
    ]


def build_manual_repository_charts(manual_repo: dict) -> dict:
    topic_rows = [[topic, count] for topic, count in manual_repo.get("by_topic", {}).items()]
    role_rows = [[role, count] for role, count in manual_repo.get("by_role", {}).items()]
    folder_rows = [[folder, count] for folder, count in manual_repo.get("by_folder_top", {}).items()]
    return {
        "topics": {"title": "Full manual repository by evidence topic", "columns": ["Topic", "Files"], "rows": topic_rows},
        "roles": {"title": "Full manual repository by evidence role", "columns": ["Role", "Files"], "rows": role_rows},
        "folders": {"title": "Largest national manual folders", "columns": ["Folder", "Files"], "rows": folder_rows},
    }


def build_it_is_charts() -> dict:
    road_rows = ITIS_TABLES["road_network_by_category"]["rows"]
    condition_rows = ITIS_TABLES["road_condition_by_category"]["rows"]
    crash_rows = ITIS_TABLES["road_crashes_by_nature"]["rows"]
    ducar_km = sum(row[1] for row in road_rows if row[0] != "National Roads")
    national_km = next(row[1] for row in road_rows if row[0] == "National Roads")
    poor_total = sum(row[3] for row in condition_rows)
    total_condition = sum(row[4] for row in condition_rows)
    crash_growth = round(((crash_rows[-1][4] - crash_rows[0][4]) / crash_rows[0][4]) * 100, 1)
    return {
        "kpis": [
            {"label": "Total road network", "value": f"{sum(row[1] for row in road_rows):,.0f} km", "note": "ITIS/URF FY 2022/23"},
            {"label": "DUCAR focus", "value": f"{ducar_km:,.0f} km", "note": "Non-national roads"},
            {"label": "National reference", "value": f"{national_km:,.0f} km", "note": "Excluded from DUCAR allocation unless delegated"},
            {"label": "Roads in poor condition", "value": f"{round((poor_total / total_condition) * 100)}%", "note": f"{poor_total:,.0f} km reported poor"},
            {"label": "Crash growth", "value": f"{crash_growth}%", "note": "CY 2019 to CY 2023"},
        ],
        "charts": {
            "network": ITIS_TABLES["road_network_by_category"],
            "condition": ITIS_TABLES["road_condition_by_category"],
            "crashes": ITIS_TABLES["road_crashes_by_nature"],
            "pavedTrend": ITIS_TABLES["paved_national_roads_trend"],
            "multimodal": ITIS_TABLES["air_passenger_and_cargo"],
            "ferry": ITIS_TABLES["ferry_operations_kis"],
        },
    }


def build_online_matrix(records: list[dict]) -> dict:
    rows = []
    group_scores: defaultdict[str, list[int]] = defaultdict(list)
    for record in records:
        score = int(record.get("evidence_score", 0))
        group_scores[record["group"]].append(score)
        rows.append([
            record["group"],
            record["title"],
            record["agency"],
            record.get("status", ""),
            record.get("words", 0),
            record.get("dominant_topic", ""),
            score,
        ])
    group_rows = [
        [group, len(scores), round(statistics.mean(scores), 1) if scores else 0]
        for group, scores in sorted(group_scores.items())
    ]
    return {
        "sourceTable": {
            "title": "Online national and global source read checks",
            "columns": ["Source group", "Title", "Agency", "Status", "Words read", "Dominant topic", "Evidence score"],
            "rows": rows,
        },
        "groupChart": {
            "title": "Online source group evidence score",
            "columns": ["Source group", "Sources", "Average evidence score"],
            "rows": group_rows,
        },
    }


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_DATA.mkdir(parents=True, exist_ok=True)

    core_paths = iter_core_documents()
    documents = [read_doc(path) for path in core_paths]
    manual_repo = read_manual_repository_catalog()
    mowt_catalog = load_json(MOWT_CATALOG, {"records": []})
    source_coverage = build_source_coverage(documents)
    tabular_extracts = build_tabular_extracts(documents)
    case_package_tables = read_case_package_tables()
    global_case_charts = build_global_case_charts(case_package_tables)
    transport_charts = build_transport_vehicle_charts()
    document_topic_chart = build_document_topic_chart(documents)

    online_records = []
    with ThreadPoolExecutor(max_workers=6) as pool:
        futures = [pool.submit(fetch_online_source, source) for source in ONLINE_SOURCES]
        for future in as_completed(futures):
            online_records.append(future.result())
    online_records.sort(key=lambda item: (item["group"], item["title"]))

    doc_pages = sum(doc.get("pages", 0) for doc in documents)
    doc_words = sum(doc.get("words", 0) for doc in documents)
    readable_docs = sum(1 for doc in documents if doc.get("status") == "read")
    payload = {
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "method": {
            "local_documents": "PDF and DOCX text extraction using pypdf and python-docx; JSON digest parsed as text.",
            "manual_repository": "Full national manual repository classified from manuals_catalog.json; unsupported legacy/binary files are used as metadata evidence.",
            "online_sources": "National and global URLs fetched with source status, text sample extraction and decision-topic scoring.",
            "apa_rule": "Every generated table carries a source string; source links remain in the Sources tab.",
        },
        "summary": {
            "core_documents_found": len(core_paths),
            "core_documents_read": readable_docs,
            "core_pages_read": doc_pages,
            "core_words_read": doc_words,
            "docx_tables_read": sum(doc.get("tables", 0) for doc in documents if doc.get("extension") == ".docx"),
            "local_tables_read": sum(doc.get("tables", 0) for doc in documents),
            "tabular_extracts": len(tabular_extracts),
            "global_case_records": (case_package_tables.get("countryCaseStudies") or {}).get("row_count", 0),
            "decision_assumptions": (case_package_tables.get("decisionAssumptions") or {}).get("row_count", 0),
            "transport_vehicle_classes": len((transport_charts.get("vehicleClassTotals") or {}).get("rows", [])),
            "manual_repository_files": manual_repo.get("file_count", 0),
            "manual_logic_records": manual_repo.get("logic_record_count", 0),
            "mowt_pdf_records": len(mowt_catalog.get("records", [])),
            "online_sources_checked": len(online_records),
            "online_sources_read": sum(1 for item in online_records if item.get("status") == "read"),
        },
        "documents": documents,
        "documentTable": build_document_table(documents),
        "documentTopicChart": document_topic_chart,
        "sourceCoverage": source_coverage,
        "tabularExtracts": tabular_extracts,
        "casePackageTables": case_package_tables,
        "globalCaseStudyCharts": global_case_charts,
        "transportCharts": transport_charts,
        "storyCards": build_story_cards(
            {
                "core_documents_found": len(core_paths),
                "core_documents_read": readable_docs,
                "manual_repository_files": manual_repo.get("file_count", 0),
                "manual_logic_records": manual_repo.get("logic_record_count", 0),
            },
            document_topic_chart,
            case_package_tables,
            transport_charts,
            source_coverage,
        ),
        "manualRepository": manual_repo,
        "manualRepositoryCharts": build_manual_repository_charts(manual_repo),
        "mowtManuals": mowt_catalog.get("records", []),
        "itisTables": ITIS_TABLES,
        "itisCharts": build_it_is_charts(),
        "onlineSources": online_records,
        "onlineEvidence": build_online_matrix(online_records),
    }
    OUT_PUBLIC.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    OUT_DATA.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    print(json.dumps(payload["summary"], indent=2))


if __name__ == "__main__":
    main()
